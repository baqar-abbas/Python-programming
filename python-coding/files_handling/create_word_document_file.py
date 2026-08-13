from docx import Document

# Create a new Word document
doc = Document()

# Add content
doc.add_heading('My First Word Document', level=1)
doc.add_paragraph("Hello from Python!")
doc.add_paragraph("This document was created using the python-docx library.")
doc.add_paragraph("You can write text, headings, and paragraphs in a .docx file.")

# Save the document
doc.save('my_word_document.docx')

print("Word document created: my_word_document.docx")